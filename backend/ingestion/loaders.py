import os
import re
import pandas as pd
from typing import Dict, Any, Tuple
import PyPDF2
import pdfplumber
from docx import Document as DocxDocument

class BaseLoader:
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.filename = os.path.basename(file_path)

    def load(self) -> Tuple[str, Dict[str, Any]]:
        """Loads file, returning a tuple of (clean_text, metadata)."""
        raise NotImplementedError("Loaders must implement load() method.")

    def clean_text(self, text: str) -> str:
        """Standardizes whitespaces, handles line breaks and strips noise."""
        if not text:
            return ""
        # Replace multiple spaces with a single space
        text = re.sub(r"[ \t]+", " ", text)
        # Normalize double linebreaks to prevent infinite gaps
        text = re.sub(r"\n\s*\n+", "\n\n", text)
        return text.strip()

    def detect_language(self, text: str) -> str:
        """Lightweight heuristic language detector based on word frequency."""
        if not text:
            return "en"
        text_lower = text[:5000].lower()
        spanish_indicators = ["el", "la", "los", "que", "y", "en", "un", "con", "para"]
        french_indicators = ["le", "la", "les", "et", "en", "un", "une", "pour", "dans"]
        german_indicators = ["der", "die", "das", "und", "ist", "in", "zu", "den", "mit"]
        
        scores = {
            "es": sum(1 for w in spanish_indicators if f" {w} " in text_lower),
            "fr": sum(1 for w in french_indicators if f" {w} " in text_lower),
            "de": sum(1 for w in german_indicators if f" {w} " in text_lower),
            "en": 2  # Baseline default
        }
        
        detected = max(scores, key=scores.get)
        return detected if scores[detected] > 2 else "en"

class PDFLoader(BaseLoader):
    def load(self) -> Tuple[str, Dict[str, Any]]:
        text_content = []
        page_count = 0
        metadata = {
            "filename": self.filename,
            "file_type": "pdf",
            "page_count": 0,
            "title": "",
            "author": ""
        }
        
        # Method 1: Try reading with pdfplumber for high fidelity (tables and layout)
        try:
            with pdfplumber.open(self.file_path) as pdf:
                page_count = len(pdf.pages)
                metadata["page_count"] = page_count
                
                # Check for document level metadata
                if pdf.metadata:
                    metadata["title"] = pdf.metadata.get("Title", "")
                    metadata["author"] = pdf.metadata.get("Author", "")
                
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)
                    # Extract tables if present
                    tables = page.extract_tables()
                    if tables:
                        for table in tables:
                            table_str = "\n"
                            for row in table:
                                if row:
                                    table_str += " | ".join([str(cell or "").strip() for cell in row]) + "\n"
                            text_content.append(table_str)
        except Exception:
            # Method 2: Fallback to PyPDF2
            try:
                with open(self.file_path, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    page_count = len(reader.pages)
                    metadata["page_count"] = page_count
                    
                    pdf_info = reader.metadata
                    if pdf_info:
                        metadata["title"] = pdf_info.get("/Title", "")
                        metadata["author"] = pdf_info.get("/Author", "")
                        
                    for page in reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_content.append(page_text)
            except Exception as e:
                raise ValueError(f"Failed to parse PDF document {self.filename}: {str(e)}")

        full_text = self.clean_text("\n\n".join(text_content))
        metadata["language"] = self.detect_language(full_text)
        return full_text, metadata

class DocxLoader(BaseLoader):
    def load(self) -> Tuple[str, Dict[str, Any]]:
        metadata = {
            "filename": self.filename,
            "file_type": "docx",
            "title": "",
            "author": ""
        }
        try:
            doc = DocxDocument(self.file_path)
            paragraphs = []
            
            # Read paragraphs
            for p in doc.paragraphs:
                if p.text.strip():
                    paragraphs.append(p.text)
            
            # Read tables
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    paragraphs.append(" | ".join(cells))

            full_text = self.clean_text("\n\n".join(paragraphs))
            metadata["language"] = self.detect_language(full_text)
            
            # Word properties
            core_properties = doc.core_properties
            if core_properties:
                metadata["title"] = core_properties.title or ""
                metadata["author"] = core_properties.author or ""
                
            return full_text, metadata
        except Exception as e:
            raise ValueError(f"Failed to parse Word document {self.filename}: {str(e)}")

class TextLoader(BaseLoader):
    def load(self) -> Tuple[str, Dict[str, Any]]:
        metadata = {
            "filename": self.filename,
            "file_type": "txt",
            "language": "en"
        }
        try:
            with open(self.file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
            clean = self.clean_text(text)
            metadata["language"] = self.detect_language(clean)
            return clean, metadata
        except Exception as e:
            raise ValueError(f"Failed to read text file {self.filename}: {str(e)}")

class CSVLoader(BaseLoader):
    def load(self) -> Tuple[str, Dict[str, Any]]:
        metadata = {
            "filename": self.filename,
            "file_type": "csv",
            "language": "en"
        }
        try:
            df = pd.read_csv(self.file_path)
            metadata["rows_count"] = len(df)
            metadata["columns"] = list(df.columns)
            
            # Formulate structured text representations
            records = []
            for i, row in df.iterrows():
                row_str = f"Record {i + 1}: " + ", ".join([f"{col}='{val}'" for col, val in row.items() if pd.notna(val)])
                records.append(row_str)
                
            full_text = "\n".join(records)
            return full_text, metadata
        except Exception as e:
            raise ValueError(f"Failed to parse CSV file {self.filename}: {str(e)}")

class ExcelLoader(BaseLoader):
    def load(self) -> Tuple[str, Dict[str, Any]]:
        metadata = {
            "filename": self.filename,
            "file_type": "excel",
            "language": "en"
        }
        try:
            xls = pd.ExcelFile(self.file_path)
            sheet_names = xls.sheet_names
            metadata["sheets"] = sheet_names
            
            all_sheet_records = []
            for sheet in sheet_names:
                df = pd.read_excel(xls, sheet_name=sheet)
                all_sheet_records.append(f"--- Sheet: {sheet} ---")
                for i, row in df.iterrows():
                    row_str = f"Row {i + 1}: " + ", ".join([f"{col}='{val}'" for col, val in row.items() if pd.notna(val)])
                    all_sheet_records.append(row_str)
                    
            full_text = "\n".join(all_sheet_records)
            return full_text, metadata
        except Exception as e:
            raise ValueError(f"Failed to parse Excel file {self.filename}: {str(e)}")

class MarkdownLoader(BaseLoader):
    def load(self) -> Tuple[str, Dict[str, Any]]:
        metadata = {
            "filename": self.filename,
            "file_type": "markdown",
            "language": "en"
        }
        try:
            with open(self.file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
                
            # Document headers mapping
            headers = [line.strip("# ").strip() for line in text.split("\n") if line.startswith("#")]
            if headers:
                metadata["title"] = headers[0]
                metadata["headers"] = headers
                
            clean = self.clean_text(text)
            metadata["language"] = self.detect_language(clean)
            return clean, metadata
        except Exception as e:
            raise ValueError(f"Failed to parse Markdown file {self.filename}: {str(e)}")

class HTMLLoader(BaseLoader):
    def load(self) -> Tuple[str, Dict[str, Any]]:
        metadata = {
            "filename": self.filename,
            "file_type": "html",
            "language": "en"
        }
        try:
            with open(self.file_path, "r", encoding="utf-8", errors="ignore") as f:
                html = f.read()
            
            # Simple BeautifulSoup-free extraction using regex to remain robust and dependency-light
            # Remove scripts and style sections
            html_clean = re.sub(r"<script.*?>.*?</script>", "", html, flags=re.DOTALL)
            html_clean = re.sub(r"<style.*?>.*?</style>", "", html_clean, flags=re.DOTALL)
            # Remove all HTML tags
            text = re.sub(r"<.*?>", " ", html_clean)
            # Unescape basic HTML entities
            text = text.replace("&nbsp;", " ").replace("&lt;", "<").replace("&gt;", ">").replace("&amp;", "&")
            
            clean = self.clean_text(text)
            metadata["language"] = self.detect_language(clean)
            
            # Try to grab title tag
            title_match = re.search(r"<title>(.*?)</title>", html, re.IGNORECASE)
            if title_match:
                metadata["title"] = title_match.group(1).strip()
                
            return clean, metadata
        except Exception as e:
            raise ValueError(f"Failed to parse HTML file {self.filename}: {str(e)}")

def get_loader(file_path: str) -> BaseLoader:
    ext = file_path.split(".")[-1].lower()
    if ext == "pdf":
        return PDFLoader(file_path)
    elif ext in ["docx"]:
        return DocxLoader(file_path)
    elif ext in ["csv"]:
        return CSVLoader(file_path)
    elif ext in ["xlsx", "xls"]:
        return ExcelLoader(file_path)
    elif ext in ["md"]:
        return MarkdownLoader(file_path)
    elif ext in ["html", "htm"]:
        return HTMLLoader(file_path)
    else:
        return TextLoader(file_path)
